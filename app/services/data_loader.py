from google.cloud import documentai_v1 as documentai
<<<<<<< HEAD
<<<<<<< HEAD
from google.cloud.documentai_v1.types import Document
from typing import List, Dict
from docx import Document as DocxDocument
from app.core.config import (
    GOOGLE_PROJECT_ID, GOOGLE_LOCATION, GOOGLE_PROCESSOR_ID,
    MIN_SECTION_TEXT_LENGTH, DEFAULT_HEADER_TEXT, CHUNK_SIZE, CHUNK_OVERLAP,
    DOCUMENT_AI_HEADER_TYPES, DOCUMENT_AI_PARAGRAPH_TYPES,
    DOCUMENT_AI_LIST_ITEM_TYPES, DOCUMENT_AI_FOOTER_TYPES, DOCUMENT_AI_TABLE_TYPES
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from app.services.llm_client import OllamaClient
from app.core.prompts import get_question_generation_prompt
import logging
import re
import os
from PyPDF2 import PdfReader
from docx import Document


=======
=======
from google.cloud.documentai_v1.types import Document
>>>>>>> e1d6f52 (Improved chunking technique)
from typing import List, Dict
from app.core.config import (
    GOOGLE_PROJECT_ID, GOOGLE_LOCATION, GOOGLE_PROCESSOR_ID,
    MIN_SECTION_TEXT_LENGTH, DEFAULT_HEADER_TEXT, CHUNK_SIZE, CHUNK_OVERLAP,
    DOCUMENT_AI_HEADER_TYPES, DOCUMENT_AI_PARAGRAPH_TYPES,
    DOCUMENT_AI_LIST_ITEM_TYPES, DOCUMENT_AI_FOOTER_TYPES, DOCUMENT_AI_TABLE_TYPES
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
import logging
import re
<<<<<<< HEAD
>>>>>>> 4499d3e (The initial version of the RAG is running smoothly)
=======
import os
>>>>>>> e1d6f52 (Improved chunking technique)

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        logger.info("--- DocumentProcessor: Initializing... ---")
        if not all([GOOGLE_PROJECT_ID, GOOGLE_LOCATION, GOOGLE_PROCESSOR_ID]):
            logger.error("--- DocumentProcessor: Google Document AI credentials (PROJECT_ID, LOCATION, or PROCESSOR_ID) are missing. ---")
            raise Exception("Google Document AI credentials not configured")

        self.project_id = GOOGLE_PROJECT_ID
        self.location = GOOGLE_LOCATION
        self.processor_id = GOOGLE_PROCESSOR_ID
        
        try:
            self.client = documentai.DocumentProcessorServiceClient()
            logger.info("--- DocumentProcessor: DocumentProcessorServiceClient initialized successfully. ---")
        except Exception as e:
            logger.error(f"--- DocumentProcessor: Failed to initialize DocumentProcessorServiceClient: {str(e)} ---", exc_info=True)
            raise Exception(f"Failed to initialize DocumentProcessorServiceClient: {str(e)}")
<<<<<<< HEAD
<<<<<<< HEAD
=======
>>>>>>> e1d6f52 (Improved chunking technique)
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len
        )
<<<<<<< HEAD

        # Initialize LLM client for question generation
        try:
            self.llm_client = OllamaClient()
            logger.info("--- DocumentProcessor: LLM client initialized for question generation. ---")
        except Exception as e:
            logger.warning(f"--- DocumentProcessor: Failed to initialize LLM client: {str(e)}. Questions will be empty. ---")
            self.llm_client = None

=======
>>>>>>> 4499d3e (The initial version of the RAG is running smoothly)
=======
>>>>>>> e1d6f52 (Improved chunking technique)
        logger.info("--- DocumentProcessor: Initialized successfully. ---")

    def process_pdf(self, file_path: str) -> List[Dict[str, str]]:
        logger.info(f"--- DocumentProcessor: Starting process_pdf for: {file_path} ---")
<<<<<<< HEAD
<<<<<<< HEAD
        document_name = os.path.basename(file_path)
=======
>>>>>>> 4499d3e (The initial version of the RAG is running smoothly)
=======
        document_name = os.path.basename(file_path)
>>>>>>> e1d6f52 (Improved chunking technique)
        try:
            name = f"projects/{self.project_id}/locations/{self.location}/processors/{self.processor_id}"
            logger.info(f"--- DocumentProcessor: Document AI Processor name: {name} ---")

            with open(file_path, "rb") as file:
                file_content = file.read()
            logger.info(f"--- DocumentProcessor: Read {len(file_content)} bytes from file: {file_path} ---")

<<<<<<< HEAD
<<<<<<< HEAD
            raw_document = documentai.RawDocument(
                content=file_content,
                mime_type="application/pdf"
            )
            
=======
            # Construct the RawDocument object
=======
>>>>>>> e1d6f52 (Improved chunking technique)
            raw_document = documentai.RawDocument(
                content=file_content,
                mime_type="application/pdf"
            )
            
<<<<<<< HEAD
            # Construct the request
>>>>>>> 4499d3e (The initial version of the RAG is running smoothly)
=======
>>>>>>> e1d6f52 (Improved chunking technique)
            request = documentai.ProcessRequest(
                name=name,
                raw_document=raw_document
            )
            logger.info("--- DocumentProcessor: Prepared Document AI request. ---")

<<<<<<< HEAD
<<<<<<< HEAD
            result = self.client.process_document(request=request)
            logger.info("--- DocumentProcessor: Successfully called Google Document AI process_document. ---")
            
            document_proto_obj: Document = result.document # Type hint with the imported Document

            if not document_proto_obj:
                logger.warning(f"--- DocumentProcessor: Document AI result for {file_path} contains no document object. ---")
                return []

            # Check if we have document_layout with blocks
            if not (hasattr(document_proto_obj, 'document_layout') and 
                    document_proto_obj.document_layout and 
                    hasattr(document_proto_obj.document_layout, 'blocks') and 
                    document_proto_obj.document_layout.blocks):
                logger.warning(f"--- DocumentProcessor: Document AI result for {file_path} contains no document_layout or blocks. ---")
                return []

            logger.info(f"--- DocumentProcessor: Found {len(document_proto_obj.document_layout.blocks)} blocks in Document AI response for {file_path}. ---")

            logger.info(f"--- These are the results from Document AI (full text snippet):\n\n{document_proto_obj.text[:500]}... ---")

            chunks = self._extract_chunks_from_layout_parser(document_proto_obj, document_name)
=======
            # This is the actual call to Google Document AI
=======
>>>>>>> e1d6f52 (Improved chunking technique)
            result = self.client.process_document(request=request)
            logger.info("--- DocumentProcessor: Successfully called Google Document AI process_document. ---")
            
            document_proto_obj: Document = result.document # Type hint with the imported Document

            if not document_proto_obj:
                logger.warning(f"--- DocumentProcessor: Document AI result for {file_path} contains no document object. ---")
                return []

<<<<<<< HEAD
            chunks = self._extract_chunks(result.document)
>>>>>>> 4499d3e (The initial version of the RAG is running smoothly)
=======
            # Check if we have document_layout with blocks
            if not (hasattr(document_proto_obj, 'document_layout') and 
                    document_proto_obj.document_layout and 
                    hasattr(document_proto_obj.document_layout, 'blocks') and 
                    document_proto_obj.document_layout.blocks):
                logger.warning(f"--- DocumentProcessor: Document AI result for {file_path} contains no document_layout or blocks. ---")
                return []

            logger.info(f"--- DocumentProcessor: Found {len(document_proto_obj.document_layout.blocks)} blocks in Document AI response for {file_path}. ---")

            logger.info(f"--- These are the results from Document AI (full text snippet):\n\n{document_proto_obj.text[:500]}... ---")

            chunks = self._extract_chunks_from_layout_parser(document_proto_obj, document_name)
>>>>>>> e1d6f52 (Improved chunking technique)
            logger.info(f"--- DocumentProcessor: Extracted {len(chunks)} chunks from Document AI response for {file_path}. ---")
            logger.info(f"--- DocumentProcessor: Finished process_pdf for: {file_path} ---")
            return chunks

        except Exception as e:
            logger.error(f"--- DocumentProcessor: Error in process_pdf for {file_path}: {str(e)} ---", exc_info=True)
<<<<<<< HEAD
<<<<<<< HEAD
            raise Exception(f"Document processing failed within DocumentProcessor: {str(e)}")


    def _process_section_into_chunks(self, header_text: str, section_content_str: str,
                                page_number: int, document_name: str,
                                previous_chunks: List[Dict[str,str]]) -> List[Dict[str, str]]:
        new_chunks_for_section = []
    
        content_to_evaluate = section_content_str
        if header_text != DEFAULT_HEADER_TEXT and section_content_str.startswith(header_text):
            content_to_evaluate = section_content_str[len(header_text):].strip()

        if not content_to_evaluate or len(content_to_evaluate) < MIN_SECTION_TEXT_LENGTH:
            if previous_chunks and previous_chunks[-1]['page'] == str(page_number) and previous_chunks[-1]['header'] == header_text:
                logger.debug(f"--- DocumentProcessor: Section under '{header_text}' too short. Merging with previous chunk. ---")
                previous_chunks[-1]['content'] += "\n" + section_content_str.strip()
                # Regenerate questions for merged content
                try:
                    questions = self._generate_questions_for_chunk(previous_chunks[-1]['content'])
                    previous_chunks[-1]['questions'] = questions
                except Exception as e:
                    logger.error(f"--- DocumentProcessor: Error generating questions for merged chunk: {str(e)} ---")
                    previous_chunks[-1]['questions'] = []
                
                if len(previous_chunks[-1]['content']) > CHUNK_SIZE:
                    logger.warning(f"--- DocumentProcessor: Merged chunk became too large ({len(previous_chunks[-1]['content'])} chars). Consider re-splitting. ---")
            elif header_text != DEFAULT_HEADER_TEXT:
                logger.debug(f"--- DocumentProcessor: Section under '{header_text}' is short but has a header. Keeping as one chunk. ---")
                try:
                    questions = self._generate_questions_for_chunk(section_content_str.strip())
                except Exception as e:
                    logger.error(f"--- DocumentProcessor: Error generating questions for short section: {str(e)} ---")
                    questions = []
                    
                new_chunks_for_section.append({
                    "content": section_content_str.strip(),
                    "page": str(page_number),
                    "header": header_text,
                    "document_name": document_name,
                    "questions": questions
                })
            else:
                logger.debug(f"--- DocumentProcessor: Section under '{header_text}' too short and no distinct header. Discarding. ---")
            return new_chunks_for_section

        if len(section_content_str) > CHUNK_SIZE:
            logger.debug(f"--- DocumentProcessor: Section under '{header_text}' (len: {len(section_content_str)}) > CHUNK_SIZE ({CHUNK_SIZE}). Splitting. ---")
            split_texts = self.text_splitter.split_text(section_content_str)
            for i, part in enumerate(split_texts):
                try:
                    questions = self._generate_questions_for_chunk(part.strip())
                except Exception as e:
                    logger.error(f"--- DocumentProcessor: Error generating questions for split chunk {i+1}: {str(e)} ---")
                    questions = []
                    
                new_chunks_for_section.append({
                    "content": part.strip(),
                    "page": str(page_number),
                    "header": header_text,
                    "document_name": document_name,
                    "questions": questions
                })
                logger.debug(f"--- DocumentProcessor: Added split chunk {i+1} for header '{header_text}'. Length: {len(part.strip())}, Questions: {len(questions)} ---")
        else:
            try:
                questions = self._generate_questions_for_chunk(section_content_str.strip())
            except Exception as e:
                logger.error(f"--- DocumentProcessor: Error generating questions for single chunk: {str(e)} ---")
                questions = []
                
            new_chunks_for_section.append({
                "content": section_content_str.strip(),
                "page": str(page_number),
                "header": header_text,
                "document_name": document_name,
                "questions": questions
            })
            logger.debug(f"--- DocumentProcessor: Added section as one chunk for header '{header_text}'. Length: {len(section_content_str.strip())}, Questions: {len(questions)} ---")
    
        return new_chunks_for_section

    def _get_text_from_nested_blocks(self, nested_blocks_list: List['Document.Layout.Block']) -> str:
        """Helper to extract and join text from a list of nested blocks."""
        parts = []
        if not nested_blocks_list:
            return ""

        for block_data in nested_blocks_list: # block_data is Document.Layout.Block
            if hasattr(block_data, 'text_block') and block_data.text_block and hasattr(block_data.text_block, 'text') and block_data.text_block.text:
                block_text = block_data.text_block.text.strip()
                if block_text:
                    parts.append(block_text)
                    if hasattr(block_data.text_block, 'blocks') and block_data.text_block.blocks: 
                        parts.append(self._get_text_from_nested_blocks(block_data.text_block.blocks))
        return "\n".join(parts)

    def _extract_chunks_from_layout_parser(self, document_proto_obj: Document, document_name: str) -> List[Dict[str, str]]:
        logger.debug(f"--- DocumentProcessor: Starting _extract_chunks_from_layout_parser for {document_name}. ---")
        all_chunks = []
        
        doc_layout: 'Document.DocumentLayout' = document_proto_obj.document_layout  # Type hint
        if not doc_layout or not doc_layout.blocks:
            logger.warning("--- DocumentProcessor: No document_layout or blocks found in Document AI response. ---")
            return all_chunks

        current_header_text = DEFAULT_HEADER_TEXT
        current_section_content_parts = []
        current_section_page_start = 1

        for block_data in doc_layout.blocks:  # block_data is Document.Layout.Block
            text_block: 'Document.Layout.TextBlock' = block_data.text_block  # Type hint
            block_type = text_block.type
            block_text_header_only = text_block.text.strip() if text_block.text else ""
            # Get page number from pageSpan
            page_start = block_data.page_span.page_start if hasattr(block_data, 'page_span') and block_data.page_span and block_data.page_span.page_start else 1

            if not block_text_header_only and not (hasattr(text_block, 'blocks') and text_block.blocks):
                continue
            
            logger.debug(f"--- DocumentProcessor: Processing block_id: {block_data.block_id}, type: {block_type}, page: {page_start}, text_header_only: '{block_text_header_only[:50]}...' ---")

            if block_type in DOCUMENT_AI_FOOTER_TYPES:
                logger.debug(f"--- DocumentProcessor: Skipping footer block. ---")
                continue
            
            if block_type in DOCUMENT_AI_TABLE_TYPES:
                logger.debug(f"--- DocumentProcessor: Skipping table block (further processing TBD). ---")
                continue

            is_new_header = block_type in DOCUMENT_AI_HEADER_TYPES

            if is_new_header:
                if current_section_content_parts:
                    section_text_str = "\n".join(current_section_content_parts).strip()
                    if section_text_str:
                        processed_chunks = self._process_section_into_chunks(
                            current_header_text, section_text_str, current_section_page_start,
                            document_name, all_chunks
                        )
                        all_chunks.extend(processed_chunks)
                
                current_header_text = block_text_header_only
                current_section_content_parts = [block_text_header_only]
                current_section_page_start = page_start
                logger.debug(f"--- DocumentProcessor: New header found: '{current_header_text}', page: {page_start}. ---")

                if hasattr(text_block, 'blocks') and text_block.blocks:
                    nested_text = self._get_text_from_nested_blocks(text_block.blocks)
                    if nested_text:
                        current_section_content_parts.append(nested_text)
                        logger.debug(f"--- DocumentProcessor: Appended nested block text to header '{current_header_text}'. ---")
            
            else:
                if block_text_header_only:
                    if not current_section_content_parts:
                        current_section_page_start = page_start
                    current_section_content_parts.append(block_text_header_only)
                    logger.debug(f"--- DocumentProcessor: Appending text to section '{current_header_text}'. ---")
                
                if hasattr(text_block, 'blocks') and text_block.blocks:
                    nested_text = self._get_text_from_nested_blocks(text_block.blocks)
                    if nested_text:
                        if not current_section_content_parts:
                            current_section_page_start = page_start
                        current_section_content_parts.append(nested_text)
                        logger.debug(f"--- DocumentProcessor: Appended nested block text to non-header section '{current_header_text}'. ---")
        
        if current_section_content_parts:
            section_text_str = "\n".join(current_section_content_parts).strip()
            if section_text_str:
                processed_chunks = self._process_section_into_chunks(
                    current_header_text, section_text_str, current_section_page_start,
                    document_name, all_chunks
                )
                all_chunks.extend(processed_chunks)

        logger.info(f"--- DocumentProcessor: Finished _extract_chunks_from_layout_parser, found {len(all_chunks)} chunks. ---")
        return all_chunks

    def _generate_questions_for_chunk(self, content: str) -> List[str]:
        """Generate common and uncommon questions for a chunk using LLM"""
        try:
            logger.debug("--- DocumentProcessor: Generating questions for chunk... ---")
            
            # Check if LLM client is available
            if not self.llm_client:
                logger.warning("--- DocumentProcessor: LLM client not available. Returning empty questions. ---")
                return []
            
            # Skip question generation for very short content
            if len(content.strip()) < 50:
                logger.debug("--- DocumentProcessor: Content too short for question generation. ---")
                return []
            
            # Detect language (simple detection)
            spanish_words = ['el', 'la', 'es', 'de', 'que', 'y', 'en', 'un', 'una', 'con', 'por', 'para']
            words = re.findall(r'\b\w+\b', content.lower())
            if not words:
                return []
                
            spanish_count = sum(1 for word in words if word in spanish_words)
            language = "spanish" if spanish_count > len(words) * 0.15 else "english"
            
            # Get prompt template
            prompt_template = get_question_generation_prompt(language)
            # Limit content to avoid token limits while preserving context
            content_for_prompt = content[:1500] if len(content) > 1500 else content
            prompt = prompt_template.format(content=content_for_prompt)
            
            # Generate questions
            response = self.llm_client.generate_response(prompt)
            
            # Parse questions from response
            questions = self._parse_questions_from_response(response)
            logger.debug(f"--- DocumentProcessor: Generated {len(questions)} questions for chunk. ---")
            
            return questions
            
        except Exception as e:
            logger.error(f"--- DocumentProcessor: Error generating questions: {str(e)} ---", exc_info=True)
            return []

    def _parse_questions_from_response(self, response: str) -> List[str]:
        """Parse questions from LLM response"""
        try:
            questions = []
            lines = response.split('\n')
            
            for line in lines:
                line = line.strip()
                # Look for numbered questions (1., 2., etc.)
                if re.match(r'^\d+\.\s*', line):
                    question = re.sub(r'^\d+\.\s*', '', line).strip()
                    if question and not question.startswith('[') and not question.endswith(']'):
                        questions.append(question)
            
            return questions[:10]  # Return max 10 questions
            
        except Exception as e:
            logger.error(f"--- DocumentProcessor: Error parsing questions: {str(e)} ---")
            return []
    
    
    # _get_text_from_layout is primarily for OCR processor output.
    def _get_text_from_layout(self, layout_proto: 'Document.Page.Layout', full_text: str) -> str:
        logger.debug(f"--- DocumentProcessor: Entering _get_text_from_layout. ---")
=======
            # Re-raise the exception so RAGService.process_document's try-except can catch it
            # and return False, or handle it as appropriate.
=======
>>>>>>> e1d6f52 (Improved chunking technique)
            raise Exception(f"Document processing failed within DocumentProcessor: {str(e)}")


    def _process_section_into_chunks(self, header_text: str, section_content_str: str, 
                                     page_number: int, document_name: str, 
                                     previous_chunks: List[Dict[str,str]]) -> List[Dict[str, str]]:
        new_chunks_for_section = []
        
        content_to_evaluate = section_content_str
        if header_text != DEFAULT_HEADER_TEXT and section_content_str.startswith(header_text):
            content_to_evaluate = section_content_str[len(header_text):].strip()

        if not content_to_evaluate or len(content_to_evaluate) < MIN_SECTION_TEXT_LENGTH:
            if previous_chunks and previous_chunks[-1]['page'] == str(page_number) and previous_chunks[-1]['header'] == header_text:
                logger.debug(f"--- DocumentProcessor: Section under '{header_text}' too short. Merging with previous chunk. ---")
                previous_chunks[-1]['content'] += "\n" + section_content_str.strip()
                if len(previous_chunks[-1]['content']) > CHUNK_SIZE:
                    logger.warning(f"--- DocumentProcessor: Merged chunk became too large ({len(previous_chunks[-1]['content'])} chars). Consider re-splitting. ---")
            elif header_text != DEFAULT_HEADER_TEXT:
                 logger.debug(f"--- DocumentProcessor: Section under '{header_text}' is short but has a header. Keeping as one chunk. ---")
                 new_chunks_for_section.append({
                    "content": section_content_str.strip(),
                    "page": str(page_number),
                    "header": header_text,
                    "document_name": document_name
                })
            else:
                logger.debug(f"--- DocumentProcessor: Section under '{header_text}' too short and no distinct header. Discarding. ---")
            return new_chunks_for_section

        if len(section_content_str) > CHUNK_SIZE:
            logger.debug(f"--- DocumentProcessor: Section under '{header_text}' (len: {len(section_content_str)}) > CHUNK_SIZE ({CHUNK_SIZE}). Splitting. ---")
            split_texts = self.text_splitter.split_text(section_content_str)
            for i, part in enumerate(split_texts):
                new_chunks_for_section.append({
                    "content": part.strip(),
                    "page": str(page_number),
                    "header": header_text,
                    "document_name": document_name
                })
                logger.debug(f"--- DocumentProcessor: Added split chunk {i+1} for header '{header_text}'. Length: {len(part.strip())} ---")
        else:
            new_chunks_for_section.append({
                "content": section_content_str.strip(),
                "page": str(page_number),
                "header": header_text,
                "document_name": document_name
            })
            logger.debug(f"--- DocumentProcessor: Added section as one chunk for header '{header_text}'. Length: {len(section_content_str.strip())} ---")
        
        return new_chunks_for_section


<<<<<<< HEAD
    def _get_text_from_layout(self, layout_proto: documentai.Document.Page.Layout, full_text: str) -> str:
        logger.debug(f"--- DocumentProcessor: Entering _get_text_from_layout ---")
>>>>>>> 4499d3e (The initial version of the RAG is running smoothly)
=======
    def _get_text_from_nested_blocks(self, nested_blocks_list: List['Document.Layout.Block']) -> str:
        """Helper to extract and join text from a list of nested blocks."""
        parts = []
        if not nested_blocks_list:
            return ""

        for block_data in nested_blocks_list: # block_data is Document.Layout.Block
            if hasattr(block_data, 'text_block') and block_data.text_block and hasattr(block_data.text_block, 'text') and block_data.text_block.text:
                block_text = block_data.text_block.text.strip()
                if block_text:
                    parts.append(block_text)
                    if hasattr(block_data.text_block, 'blocks') and block_data.text_block.blocks: 
                        parts.append(self._get_text_from_nested_blocks(block_data.text_block.blocks))
        return "\n".join(parts)

    def _extract_chunks_from_layout_parser(self, document_proto_obj: Document, document_name: str) -> List[Dict[str, str]]:
        logger.debug(f"--- DocumentProcessor: Starting _extract_chunks_from_layout_parser for {document_name}. ---")
        all_chunks = []
        
        doc_layout: 'Document.DocumentLayout' = document_proto_obj.document_layout # Type hint
        if not doc_layout or not doc_layout.blocks:
            logger.warning("--- DocumentProcessor: No document_layout or blocks found in Document AI response. ---")
            return all_chunks

        current_header_text = DEFAULT_HEADER_TEXT
        current_section_content_parts = []
        current_section_page_start = 1

        for block_data in doc_layout.blocks: # block_data is Document.Layout.Block
            text_block: 'Document.Layout.TextBlock' = block_data.text_block # Type hint
            block_type = text_block.type
            block_text_header_only = text_block.text.strip() if text_block.text else ""
            page_start = block_data.page_span.page_start if hasattr(block_data, 'page_span') and block_data.page_span else 1

            if not block_text_header_only and not (hasattr(text_block, 'blocks') and text_block.blocks):
                continue
            
            logger.debug(f"--- DocumentProcessor: Processing block_id: {block_data.block_id}, type: {block_type}, page: {page_start}, text_header_only: '{block_text_header_only[:50]}...' ---")

            if block_type in DOCUMENT_AI_FOOTER_TYPES:
                logger.debug(f"--- DocumentProcessor: Skipping footer block. ---")
                continue
            
            if block_type in DOCUMENT_AI_TABLE_TYPES:
                logger.debug(f"--- DocumentProcessor: Skipping table block (further processing TBD). ---")
                continue

            is_new_header = block_type in DOCUMENT_AI_HEADER_TYPES

            if is_new_header:
                if current_section_content_parts:
                    section_text_str = "\n".join(current_section_content_parts).strip()
                    if section_text_str:
                        processed_chunks = self._process_section_into_chunks(
                            current_header_text, section_text_str, current_section_page_start, 
                            document_name, all_chunks 
                        )
                        all_chunks.extend(processed_chunks)
                
                current_header_text = block_text_header_only
                current_section_content_parts = [block_text_header_only] 
                current_section_page_start = page_start
                logger.debug(f"--- DocumentProcessor: New header found: '{current_header_text}', page: {page_start}. ---")

                if hasattr(text_block, 'blocks') and text_block.blocks:
                    nested_text = self._get_text_from_nested_blocks(text_block.blocks) 
                    if nested_text:
                        current_section_content_parts.append(nested_text)
                        logger.debug(f"--- DocumentProcessor: Appended nested block text to header '{current_header_text}'. ---")
            
            else: 
                if block_text_header_only:
                    if not current_section_content_parts:
                        current_section_page_start = page_start
                    current_section_content_parts.append(block_text_header_only)
                    logger.debug(f"--- DocumentProcessor: Appending text to section '{current_header_text}'. ---")
                
                if hasattr(text_block, 'blocks') and text_block.blocks:
                    nested_text = self._get_text_from_nested_blocks(text_block.blocks)
                    if nested_text:
                        if not current_section_content_parts:
                             current_section_page_start = page_start
                        current_section_content_parts.append(nested_text)
                        logger.debug(f"--- DocumentProcessor: Appended nested block text to non-header section '{current_header_text}'. ---")
        
        if current_section_content_parts:
            section_text_str = "\n".join(current_section_content_parts).strip()
            if section_text_str:
                processed_chunks = self._process_section_into_chunks(
                    current_header_text, section_text_str, current_section_page_start, 
                    document_name, all_chunks
                )
                all_chunks.extend(processed_chunks)

        logger.info(f"--- DocumentProcessor: Finished _extract_chunks_from_layout_parser, found {len(all_chunks)} chunks. ---")
        return all_chunks

    # _get_text_from_layout is primarily for OCR processor output.
    def _get_text_from_layout(self, layout_proto: 'Document.Page.Layout', full_text: str) -> str:
        logger.debug(f"--- DocumentProcessor: Entering _get_text_from_layout. ---")
>>>>>>> e1d6f52 (Improved chunking technique)
        try:
            if not layout_proto or not layout_proto.text_anchor or not layout_proto.text_anchor.text_segments:
                logger.debug("--- DocumentProcessor: _get_text_from_layout - Layout, text_anchor, or segments missing. ---")
                return ""

            text_content_parts = []
            for segment in layout_proto.text_anchor.text_segments:
                start_index = segment.start_index
                end_index = segment.end_index
                start_index = int(start_index) if start_index is not None else 0
<<<<<<< HEAD
<<<<<<< HEAD
                end_index = int(end_index) if end_index is not None else 0 
=======
                end_index = int(end_index) if end_index is not None else 0 # If end_index is None, make segment empty
>>>>>>> 4499d3e (The initial version of the RAG is running smoothly)
=======
                end_index = int(end_index) if end_index is not None else 0 
>>>>>>> e1d6f52 (Improved chunking technique)

                if not (0 <= start_index <= len(full_text) and 0 <= end_index <= len(full_text) and start_index <= end_index):
                    logger.warning(f"--- DocumentProcessor: Invalid text segment indices: start='{segment.start_index}', end='{segment.end_index}'. Full text length: {len(full_text)}. Skipping segment. ---")
                    continue
                
                text_content_parts.append(full_text[start_index:end_index])
            
            return "".join(text_content_parts)

        except Exception as e:
            logger.error(f"--- DocumentProcessor: Error in _get_text_from_layout: {str(e)} ---", exc_info=True)
<<<<<<< HEAD
<<<<<<< HEAD
            raise Exception(f"Text extraction from layout failed: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extracts full text from a DOCX file using python-docx."""
        logger.info(f"--- DocProcessor: Starting extract_text_from_docx for: {file_path} ---")
        if Document is None:
            raise ImportError("python-docx is not available. Cannot extract text from DOCX.")

        try:
            doc = Document(file_path)
            full_text = [paragraph.text for paragraph in doc.paragraphs]
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_cells_text = [cell.text for cell in row.cells]
                    tables_text.append("\t".join(row_cells_text)) # Join cells by tab for table readability
            
            # Combine paragraph and table text
            extracted_content = "\n".join(full_text + tables_text)
            logger.info(f"--- DocProcessor: Extracted {len(extracted_content)} characters from DOCX: {file_path} ---")
            return extracted_content

        except Exception as e:
            logger.error(f"--- DocProcessor: Error extracting text from DOCX {file_path}: {str(e)} ---", exc_info=True)
            raise Exception(f"DOCX text extraction failed: {str(e)}")
        
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extracts full text from a PDF file using PyPDF2."""
        logger.info(f"--- DocProcessor: Starting extract_text_from_pdf for: {file_path} ---")
        if PdfReader is None:
            raise ImportError("PyPDF2 is not available. Cannot extract text from PDF.")

        try:
            full_text = []
            with open(file_path, "rb") as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    full_text.append(page.extract_text() or "") # .extract_text() can return None
            
            extracted_content = "\n".join(full_text)
            logger.info(f"--- DocProcessor: Extracted {len(extracted_content)} characters from PDF: {file_path} ---")
            return extracted_content

        except Exception as e:
            logger.error(f"--- DocProcessor: Error extracting text from PDF {file_path}: {str(e)} ---", exc_info=True)
            raise Exception(f"PDF text extraction failed: {str(e)}")
=======
            # Re-raise to be caught by _extract_chunks
            raise Exception(f"Text extraction from layout failed: {str(e)}")
>>>>>>> 4499d3e (The initial version of the RAG is running smoothly)
=======
            raise Exception(f"Text extraction from layout failed: {str(e)}")
>>>>>>> e1d6f52 (Improved chunking technique)
